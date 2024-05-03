import os
import shutil
import time

from pptx import Presentation
import loguru
from docx import Document

from AI_PPT import structure_article, yiyan
from AI_PPT.mdtree import tree2ppt
from src.presets import PAPER_PROMPT_TEMPLATE, test_site_template
from src.model import ChatGLMClient
from src.presets import PPT_PROMPT_TEMPLATE, JSON_FORMAT, gist


# export DASHSCOPE_API_KEY=sk-IANIL0bbyr


# 生成PPT内容

def gen_requirement(task_name):
    Requirement_prompt = (f"""你是一个专业的需求分析师，很擅长引导人将一个想法从多角度细化。我会给你一个任务或者一个主题，请你通过反问的形式提出针对该任务或主题细化后的几个角度，帮我更加明白我自己的想法。"
                          任务名称:{task_name}
                          输出格式：
                          1.
                          2.
                          3. 
                          """)
    result = yiyan.call_with_messages(Requirement_prompt)
    return result


# 生成PPT文件
def save_knowledge_func(task_name, knowledge_content, mode, sub_num):
    time1 = time.time()
    sub_num = int(sub_num)
    character_a = "你是一个精通各方面知识的人"
    struct_article = structure_article.StructureArticle(api_type="yiyan", main_idea_knowledge=knowledge_content,
                                                        max_sub_idea_num=sub_num, min_sub_idea_num=sub_num)
    content = struct_article.generate_final_summary(task_name, character_a)
    if len(os.listdir("../ppt")) > 100:
        shutil.rmtree("../ppt")
        os.makedirs("../ppt")

    loguru.logger.debug(f'{type(mode)},{type(sub_num)}')
    save_path = "./ppt/" + task_name + str(mode) + str(sub_num) + ".pptx"
    tree2ppt.Tree2PPT(content, "AI_PPT/my_ppt_mode/" + str(mode), save_path=save_path)

    print("一共消耗时长：", time.time() - time1)

    return save_path


def generate_gist(model, subject, grade, content):
    prompts = gist.format(subject=subject, grade=grade, content=content)
    response = model.get_answer_at_once(prompts, [])
    loguru.logger.debug(response)
    # 创建一个新的Word文档对象
    doc = Document()
    doc.add_paragraph(response)
    # 保存Word文档
    path = 'gist/' + subject + grade + content + '.docx'
    doc.save(path)

    return path


def geneate_test(model, school, subject, chapter):
    prompts = test_site_template.format(school=school, subject=subject, chapter=chapter)
    response = model.get_answer_at_once(prompts, [])
    loguru.logger.debug(response)
    # 创建一个新的Word文档对象
    doc = Document()
    doc.add_paragraph(response)
    # 保存Word文档
    path = 'test/' + school + subject + chapter + '.docx'
    doc.save(path)

    return path


def generate_paper(model, topic, subject, type, number):
    # prompt
    user_prompt = """\t用户的指令：\n\t请随机给出{number}道关于{topic}的{type}, 该题目的所属科目是 {subject}""".format(
        topic=topic,
        subject=subject,
        type=type,
        number=number)
    prompt = PAPER_PROMPT_TEMPLATE + '\n' + user_prompt
    loguru.logger.info(prompt)
    response = eval(model.get_answer_at_once(prompt, []))
    loguru.logger.info(response)
    # 创建一个新的Word文档对象
    doc = Document()
    # 添加题目
    doc.add_heading('题目', level=1)
    # 遍历题目和答案列表
    for question in response:
        doc.add_paragraph(question["question"])
    doc.add_page_break()
    # 添加答案
    doc.add_heading('答案', level=1)
    for answer in response:
        doc.add_paragraph(answer["answer"])

    # 保存Word文档
    path = 'paper/' + topic + '.docx'
    doc.save(path)
    return path


if __name__ == '__main__':
    model = ChatGLMClient()
    # generate_paper(model,"有关中国外交的近代史","历史","选择题",8)
    # generate_gist(model,"历史","三年级","有关中国外交的近代史")
