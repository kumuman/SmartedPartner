# -*- coding: utf-8 -*-

import base64
import json

import loguru
import numpy as np
import cv2
import requests
from docx import Document

from utils.AuthV3Util import addAuthParams

APP_KEY = '0544f7d7c32f11d9'
APP_SECRET = 'MjImJ0RNmIkYroGEkmGcM3C8YVZv6GsU'



def vertical_concatenate_images(image_paths, output_path):
    images = []
    for img_path in image_paths:
        img = cv2.imread(img_path)
        if img is not None:
            images.append(img)

    if len(images) == 0:
        print("No images found in the list.")
        return

    # 获取所有图片的高度和最大宽度
    max_width = max(img.shape[1] for img in images)
    total_height = sum(img.shape[0] for img in images)

    # 创建一张新的空白图片
    result = np.zeros((total_height, max_width, 3), dtype=np.uint8)

    # 将每张图片拼接到新图片中
    y_offset = 0
    for img in images:
        result[y_offset:y_offset+img.shape[0], :img.shape[1]] = img
        y_offset += img.shape[0]

    # 保存结果图片
    cv2.imwrite(output_path, result)
    return output_path


def extract_keys_data(data, keys_data=None, prefix=''):
    if keys_data is None:
        keys_data = {}
    if isinstance(data, dict):
        for key, value in data.items():
            new_prefix = f'{prefix}.{key}' if prefix else key
            extract_keys_data(value, keys_data, new_prefix)
            keys_data[new_prefix] = value
    elif isinstance(data, list):
        for idx, item in enumerate(data):
            new_prefix = f'{prefix}[{idx}]' if prefix else f'[{idx}]'
            extract_keys_data(item, keys_data, new_prefix)
            keys_data[new_prefix] = item
    return keys_data

def createRequest(path, requirement):
    '''
    note: 将下列变量替换为需要请求的参数
    取值参考文档: https://ai.youdao.com/DOCSIRMA/html/%E4%BD%9C%E6%96%87%E6%89%B9%E6%94%B9/API%E6%96%87%E6%A1%A3/%E8%8B%B1%E8%AF%AD%E4%BD%9C%E6%96%87%E6%89%B9%E6%94%B9%EF%BC%88%E6%96%87%E6%9C%AC%E8%BE%93%E5%85%A5%EF%BC%89/%E8%8B%B1%E8%AF%AD%E4%BD%9C%E6%96%87%E6%89%B9%E6%94%B9%EF%BC%88%E6%96%87%E6%9C%AC%E8%BE%93%E5%85%A5%EF%BC%89-API%E6%96%87%E6%A1%A3.html
    '''
    grade = None
    # title = '作文标题'
    title = None
    # model_content = '作文参考范文'
    model_content = None
    # is_need_synonyms = '是否查询同义词'
    is_need_synonyms = None
    # correct_version = "作文批改版本：基础，高级"
    correct_version = None
    # is_need_essay_report = "是否返回写作报告"
    is_need_essay_report = None

    # requirement = "本手、妙手、俗手是围棋的三个术语。本手是指合乎棋理的正规下法；妙手是指出人意料的精妙下法；俗手是指貌似合理，而从全局看通常会受损的下法。对于初学者而言，应该从本手开始，本手的功夫扎实了，棋力才会提高。一些初学者热衷于追求妙手，而忽视更为常用的本手。一般来说，对本手理解深刻，才可能出现妙手；否则，难免下出俗手，水平也不易提升。\n以上材料对我们颇具启示意义。请结合材料写一篇文章，体现你的感悟与思考。"

    # 数据的base64编码
    q = readFileAsBase64(path)
    data = {'q': q, 'grade': grade, 'to': title, 'modelContent': model_content, 'isNeedSynonyms': is_need_synonyms,
            'correctVersion': correct_version, 'isNeedEssayReport': is_need_essay_report, 'requirement':requirement}

    addAuthParams(APP_KEY, APP_SECRET, data)

    header = {'Content-Type': 'application/x-www-form-urlencoded'}
    res = doCall('https://openapi.youdao.com/v2/correct_writing_cn_image', header, data, 'post')
    content_str = res.content.decode('utf-8')  # 将bytes转换为str
    content_dict = json.loads(content_str)  # 将str转换为字典
    loguru.logger.debug(content_dict)
    result = content_dict['Result']
    scoreCollection = result['scoreCollection']         # 所有的评分信息
    commentCollection = result['commentCollection']     # 所有的评价信息
    orgContent = result['orgContent']                   # 原始的作文内容
    detailedEvaluation = result['detailedEvaluation']
    # print(res.content)
    # print(str(res.content, 'utf-8'))
    # 提取所有键和对应数据并打印
    keys_data = extract_keys_data(result)
    # print(keys_data['scoreCollection'].get('perspectiveScore'))
    # for key, value in keys_data.items():
    #     print(f"{key}: {value}")

    # for key, value in keys_data['scoreCollection'].get('perspectiveScore').items():
    #     print(f"{key}: {value}")

    #  打分
    perspective_score = keys_data['scoreCollection']['perspectiveScore']
    score = keys_data['scoreCollection']['score']
    sentiment_sincerity = perspective_score['sentimentSincerity']
    essay_fluence = perspective_score['essayFluence']
    structure_strict = perspective_score['structureStrict']
    theme_explicit = perspective_score['themeExplicit']
    good_sent = perspective_score['goodSent']
    satisfy_requirement = perspective_score['satisfyRequirement']

    # print(f"感情真挚: {sentiment_sincerity} 分")
    # print(f"语言流畅: {essay_fluence} 分")
    # print(f"结构严谨: {structure_strict} 分")
    # # print(f"主题明确: {theme_explicit} 分")
    # print(f"好词好句: {good_sent} 分")
    # print(f"符合题意: {satisfy_requirement} 分")

    #  评价
    commentCollection = keys_data['commentCollection']
    aspectComment = commentCollection['aspectComment']

    comment = commentCollection['comment']
    emotion = aspectComment['emotion']
    langExpression = str(aspectComment['langExpression'])
    structLogic = aspectComment['structLogic']
    contentInfo = aspectComment['contentInfo']

    # print(f'整体评价: {comment}')
    # print(f'情感方面: {emotion}')
    # print(f'语言表达方面: {langExpression}')
    # print(f'结构逻辑方面: {structLogic}')
    # print(f'文章内容方面: {contentInfo}')

    #  好词好句
    detailedEvaluation = keys_data['detailedEvaluation']
    sentenceEvaluation = detailedEvaluation['sentenceEvaluation']
    phraseEvaluation = detailedEvaluation['phraseEvaluation']
    # print(f'好词好句: {sentenceEvaluation}')

    orgContent = keys_data['orgContent']
    all_comment = '整体分数: ' + str(score) + ' 分\n' + '感情真挚:' + str(sentiment_sincerity) + ' 分\n' + '语言流畅: ' \
                  + str(essay_fluence) + ' 分\n' + '结构严谨: ' + str(structure_strict) + ' 分\n' + '好词好句: ' + str(
        good_sent) + ' 分\n' + '符合题意: ' \
                  + str(
        satisfy_requirement) + ' 分\n' + '整体评价: ' + comment + '\n' + '情感方面: ' + emotion + '\n' + '语言表达方面: ' \
                  + langExpression + '\n' + '结构逻辑方面: ' + structLogic + '\n' + '文章内容方面: ' + contentInfo + '\n'

    artical = str(orgContent)

    return all_comment, artical
    # print(f'作文内容：{orgContent}')


def doCall(url, header, params, method):
    if 'get' == method:
        return requests.get(url, params)
    elif 'post' == method:
        return requests.post(url, params, header)


def readFileAsBase64(path):
    f = open(path, 'rb')
    data = f.read()
    return str(base64.b64encode(data), 'utf-8')


def get_composition(img1,img2,img3,topic,output_path):
    image_paths = [img  for img in (img1,img2,img3) if img]
    loguru.logger.debug(image_paths)
    path = vertical_concatenate_images(image_paths, output_path)
    all_comment, artical = createRequest(path, topic)
    loguru.logger.debug(all_comment)
    # 创建一个新的Word文档对象
    doc = Document()
    doc.add_paragraph(all_comment)
    doc.add_paragraph(artical)

    doc_template = Document('composition/doc_compositon_template.docx')

    for element in doc_template.element.body:
        doc.element.body.append(element)

    # 保存Word文档
    path = '../composition/essay.docx'
    doc.save(path)
    return path

# 网易有道智云英文图片作文批改服务api调用demo
# api接口: https://openapi.youdao.com/v2/correct_writing_image
if __name__ == '__main__':
    image_paths = [
        'artical1.png',
        'artical2.png',
        'artical3.png'
    ]
    output_path = '../output_img.png'
    requirement = "本手、妙手、俗手是围棋的三个术语。本手是指合乎棋理的正规下法；妙手是指出人意料的精妙下法；俗手是指貌似合理，而从全局看通常会受损的下法。对于初学者而言，应该从本手开始，本手的功夫扎实了，棋力才会提高。一些初学者热衷于追求妙手，而忽视更为常用的本手。一般来说，对本手理解深刻，才可能出现妙手；否则，难免下出俗手，水平也不易提升。\n以上材料对我们颇具启示意义。请结合材料写一篇文章，体现你的感悟与思考。"


